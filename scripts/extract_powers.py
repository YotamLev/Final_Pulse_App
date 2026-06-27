"""Extract Discipline, Amalgam, and Sorcery powers from the rulebook docx."""

from __future__ import annotations

import json
import re
import unicodedata
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "Final Pulse text - Second Edition.docx"
OUTPUT_PATH = ROOT / "data" / "powers.json"

DISCIPLINES = [
    "Animalism",
    "Auspex",
    "Celerity",
    "Dominate",
    "Fortitude",
    "Nightmare",
    "Obfuscate",
    "Potence",
    "Presence",
    "Protean",
    "Arrete",
    "Mortifex",
]

SORCERY_PATHS = [
    "Path of Koldun",
    "Path of Force",
    "Path of Wards",
    "Path of Vines",
    "Path of Shadows",
    "Path of Blood",
    "Path of Necromancy",
    "Path of Demons",
    "Path of Illusion",
]

REQ_LINE = re.compile(r"^Requirements?:\s*(.+)$", re.IGNORECASE)
PREDATOR_BONUS_LINE = re.compile(
    r"^If chosen as predator(?: type)? power:\s*(.+)$",
    re.IGNORECASE,
)

# Canonical power name fixes for prerequisite text that drifts from headings.
POWER_ALIASES = {
    "draw on instinct": "Draw On Instinct",
}


def slugify(name: str) -> str:
    normalized = unicodedata.normalize("NFKD", name)
    ascii_name = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^a-z0-9]+", "_", ascii_name.lower()).strip("_")
    return slug


def source_type_for(parent: str) -> str:
    if parent == "Amalgams":
        return "amalgam"
    if parent.startswith("Path of "):
        return "sorcery_path"
    return "discipline"


def canonical_power_name(name: str) -> str:
    key = name.strip().lower()
    return POWER_ALIASES.get(key, name.strip())


def split_power_list(text: str) -> list[str]:
    parts = re.split(r"\s*,\s*|\s+and\s+|\s*\+\s*", text, flags=re.IGNORECASE)
    return [canonical_power_name(p.strip()) for p in parts if p.strip()]


def parse_discipline_pair(text: str) -> list[str] | None:
    match = re.match(r"^([A-Za-z]+)\s*&\s*([A-Za-z]+)$", text.strip())
    if not match:
        return None
    d1, d2 = match.group(1), match.group(2)
    known = set(DISCIPLINES)
    if d1 in known and d2 in known:
        return [d1, d2]
    return None


def parse_power_options(text: str) -> list[str]:
    """Parse a comma/or-separated list of power names."""
    text = text.strip()
    if text.lower().startswith("one of:"):
        text = text.split(":", 1)[1].strip()
    parts = re.split(r"\s*,\s*|\s+or\s+", text, flags=re.IGNORECASE)
    result: list[str] = []
    for part in parts:
        cleaned = part.strip()
        if cleaned.lower().startswith("or "):
            cleaned = cleaned[3:].strip()
        if cleaned:
            result.append(canonical_power_name(cleaned))
    return result


def parse_amalgam_prerequisites(raw: str) -> dict:
    """Parse amalgam requirement lines."""
    # "Obfuscate & Potence or Obfuscate & Celerity"
    if "," not in raw:
        or_parts = re.split(r"\s+or\s+", raw, flags=re.IGNORECASE)
        if len(or_parts) > 1 and all("&" in part for part in or_parts):
            pairs = [parse_discipline_pair(part.strip()) for part in or_parts]
            if all(pairs):
                return {
                    "raw": raw,
                    "any": [{"type": "disciplines", "disciplines": pair} for pair in pairs],
                }

    segments = [segment.strip() for segment in raw.split(",") if segment.strip()]
    all_rules: list[dict] = []

    start_idx = 0
    first_pair = parse_discipline_pair(segments[0])
    if first_pair:
        all_rules.append({"type": "disciplines", "disciplines": first_pair})
        start_idx = 1

    remainder = ", ".join(segments[start_idx:]).strip()
    if remainder:
        specialty_match = re.match(r"(.+?)\s+Specialty:\s*(.+)$", remainder, re.IGNORECASE)
        if specialty_match:
            all_rules.append(
                {
                    "type": "specialty",
                    "skill": specialty_match.group(1).strip(),
                    "specialty": specialty_match.group(2).strip(),
                }
            )
        elif " or " in remainder.lower() and "&" not in remainder:
            powers = parse_power_options(remainder)
            all_rules.append({"type": "any_powers", "powers": powers})
        else:
            powers = parse_power_options(remainder)
            if len(powers) == 1:
                all_rules.append({"type": "power", "power": powers[0]})
            elif len(powers) > 1:
                all_rules.append({"type": "any_powers", "powers": powers})

    if all_rules:
        return {"raw": raw, "all": all_rules}
    return {"raw": raw, "all": []}


def parse_prerequisites(raw: str, source: str, source_type: str, power_name: str) -> dict:
    """Best-effort structured parse; always retains raw text."""
    if source_type == "amalgam":
        return parse_amalgam_prerequisites(raw)

    result: dict = {"raw": raw, "all": []}
    lower = raw.lower()

    # "N other X power(s)"
    count_match = re.match(
        r"^(\d+)\s+other\s+(" + "|".join(DISCIPLINES + SORCERY_PATHS) + r")\s+powers?$",
        raw,
        re.IGNORECASE,
    )
    if count_match:
        count = int(count_match.group(1))
        src = count_match.group(2)
        st = "sorcery_path" if src.startswith("Path of ") else "discipline"
        result["all"].append(
            {
                "type": "powers_from_source",
                "source_type": st,
                "source": src,
                "min_count": count,
                "exclude": [power_name],
            }
        )
        return result

    single_other = re.match(
        r"^1\s+other\s+(" + "|".join(DISCIPLINES + SORCERY_PATHS) + r")\s+power$",
        raw,
        re.IGNORECASE,
    )
    if single_other:
        src = single_other.group(1)
        st = "sorcery_path" if src.startswith("Path of ") else "discipline"
        result["all"].append(
            {
                "type": "powers_from_source",
                "source_type": st,
                "source": src,
                "min_count": 1,
                "exclude": [power_name],
            }
        )
        return result

    if lower.startswith("another ") and lower.endswith(" power"):
        src = raw.split()[1]
        st = "sorcery_path" if src.startswith("Path of ") else "discipline"
        result["all"].append(
            {
                "type": "powers_from_source",
                "source_type": st,
                "source": src,
                "min_count": 1,
                "exclude": [power_name],
            }
        )
        return result

    if lower.startswith("any ") and "power" in lower:
        src = source
        result["all"].append(
            {
                "type": "powers_from_source",
                "source_type": source_type,
                "source": src,
                "min_count": 1,
                "exclude": [power_name],
            }
        )
        return result

    # "Mesmerize + 2 other Dominate powers"
    plus_other = re.match(
        r"^(.+?)\s*\+\s*(\d+)\s+other\s+(" + "|".join(DISCIPLINES) + r")\s+powers?$",
        raw,
        re.IGNORECASE,
    )
    if plus_other:
        req_power = canonical_power_name(plus_other.group(1))
        count = int(plus_other.group(2))
        src = plus_other.group(3)
        result["all"].append({"type": "power", "power": req_power})
        result["all"].append(
            {
                "type": "powers_from_source",
                "source_type": "discipline",
                "source": src,
                "min_count": count,
                "exclude": [power_name],
            }
        )
        return result

    # "A or B"
    if " or " in lower and "+" not in raw:
        options = re.split(r"\s+or\s+", raw, flags=re.IGNORECASE)
        powers = [canonical_power_name(o.strip()) for o in options]
        return {"raw": raw, "any": [{"type": "power", "power": p} for p in powers]}

    # Comma / and-separated power names
    if "," in raw or re.search(r"\band\b", raw, re.IGNORECASE):
        powers = split_power_list(raw)
        if len(powers) > 1:
            result["all"] = [{"type": "power", "power": p} for p in powers]
            return result

    # Single power
    result["all"] = [{"type": "power", "power": canonical_power_name(raw)}]
    return result


def detect_predator_bonus(body_paragraphs: list[str]) -> tuple[bool, str | None]:
    """Detect predator-type bonus lines in full power body text."""
    for text in body_paragraphs:
        stripped = text.strip()
        if PREDATOR_BONUS_LINE.match(stripped):
            return True, stripped
    return False, None


def collect_unknown_power_refs(powers: list[dict]) -> list[dict]:
    """Find prerequisite power names that do not match any extracted power."""
    power_names = {entry["name"] for entry in powers}
    unknown: list[dict] = []

    def walk(power_name: str, obj: object) -> None:
        if isinstance(obj, dict):
            if obj.get("type") == "power" and obj["power"] not in power_names:
                unknown.append({"referenced_in": power_name, "power": obj["power"]})
            if obj.get("type") == "any_powers":
                for name in obj.get("powers", []):
                    if name not in power_names:
                        unknown.append({"referenced_in": power_name, "power": name})
            for value in obj.values():
                walk(power_name, value)
        elif isinstance(obj, list):
            for item in obj:
                walk(power_name, item)

    for entry in powers:
        prereq = entry.get("prerequisites")
        if prereq:
            walk(entry["name"], prereq)

    deduped: list[dict] = []
    seen: set[tuple[str, str]] = set()
    for item in unknown:
        key = (item["referenced_in"], item["power"])
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def extract_powers() -> dict:
    doc = Document(DOCX_PATH)
    current_parent: str | None = None
    powers: list[dict] = []
    seen_ids: set[str] = set()

    for i, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name
        text = paragraph.text.strip()
        if not text:
            continue

        if style == "Heading 1" and text != "Disciplines":
            if text in DISCIPLINES or text == "Amalgams" or text in SORCERY_PATHS:
                current_parent = text
            elif current_parent and i > 1485:
                current_parent = None
            continue

        if style != "Heading 2" or not current_parent:
            continue

        power_name = text
        st = source_type_for(current_parent)
        source = current_parent if st != "amalgam" else "Amalgams"

        prereq_raw: str | None = None
        summary: str | None = None
        body_paragraphs: list[str] = []

        for j in range(i + 1, len(doc.paragraphs)):
            nxt = doc.paragraphs[j]
            if nxt.style.name.startswith("Heading"):
                break
            body = nxt.text.strip()
            if not body:
                continue
            req_match = REQ_LINE.match(body)
            if req_match and prereq_raw is None:
                prereq_raw = req_match.group(1).strip()
                continue
            body_paragraphs.append(body)
            if summary is None:
                summary = body

        predator_bonus, predator_bonus_text = detect_predator_bonus(body_paragraphs)

        base_id = slugify(power_name)
        parent_slug = slugify(current_parent)
        power_id = f"{parent_slug}_{base_id}"
        if power_id in seen_ids:
            power_id = f"{parent_slug}_{base_id}_{i}"
        seen_ids.add(power_id)

        entry: dict = {
            "id": power_id,
            "name": power_name,
            "source_type": st,
            "source": source,
            "prerequisites": None,
            "predator_bonus": predator_bonus,
        }
        if summary:
            entry["summary"] = summary
        if predator_bonus_text:
            entry["predator_bonus_text"] = predator_bonus_text
        if prereq_raw:
            entry["prerequisites"] = parse_prerequisites(
                prereq_raw, source=current_parent, source_type=st, power_name=power_name
            )
        powers.append(entry)

    unresolved = collect_unknown_power_refs(powers)
    return {
        "schema_version": 1,
        "meta": {
            "source": "Final Pulse text - Second Edition",
            "source_file": DOCX_PATH.name,
            "power_count": len(powers),
            "unresolved_power_references": unresolved,
        },
        "disciplines": DISCIPLINES,
        "sorcery_paths": SORCERY_PATHS,
        "powers": powers,
    }


def main() -> None:
    data = extract_powers()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    with_prereq = sum(1 for p in data["powers"] if p.get("prerequisites"))
    with_predator = sum(1 for p in data["powers"] if p.get("predator_bonus"))
    print(
        f"Wrote {data['meta']['power_count']} powers "
        f"({with_prereq} with prerequisites, {with_predator} with predator bonus) to {OUTPUT_PATH}"
    )


if __name__ == "__main__":
    main()
